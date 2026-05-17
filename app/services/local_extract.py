from __future__ import annotations

from pathlib import Path

from app.models import ExtractedPage
from app.services.files import StoredUpload


def extract_text_locally(upload: StoredUpload) -> list[ExtractedPage]:
    if upload.mime_type == "text/plain":
        return [
            ExtractedPage(
                page_number=1,
                source_name=upload.source_name,
                language="unknown",
                text=upload.path.read_text(encoding="utf-8", errors="ignore"),
                confidence="high",
            )
        ]

    if upload.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(upload)

    if upload.mime_type == "application/pdf":
        return _extract_pdf_text(upload)

    return []


def _extract_docx(upload: StoredUpload) -> list[ExtractedPage]:
    try:
        from docx import Document
    except ImportError:
        return []

    document = Document(str(upload.path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    text = "\n".join(paragraphs).strip()
    return [
        ExtractedPage(
            page_number=1,
            source_name=upload.source_name,
            language="unknown",
            text=text,
            confidence="high" if text else "low",
            warnings=[] if text else ["No selectable text found in this Word document."],
        )
    ]


def _extract_pdf_text(upload: StoredUpload) -> list[ExtractedPage]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return []

    pages: list[ExtractedPage] = []
    try:
        reader = PdfReader(str(upload.path))
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(
                    ExtractedPage(
                        page_number=index,
                        source_name=upload.source_name,
                        language="unknown",
                        text=text,
                        confidence="medium",
                    )
                )
    except Exception:
        return []
    return pages


def is_binary_visual(path: Path, mime_type: str) -> bool:
    return mime_type.startswith("image/") or mime_type == "application/pdf"
